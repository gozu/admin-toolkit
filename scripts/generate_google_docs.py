#!/usr/bin/env python3
"""Generate Google Docs-friendly DOCX files from the internal Markdown sources.

The renderer intentionally supports only the small Markdown subset used by the
Admin Toolkit playbook and setup guide. It avoids floating shapes, text boxes,
embedded fonts, and other Word features that commonly shift during Google Docs
import. The output uses semantic heading styles, ordinary tables, paragraph
borders, standard bullets, and a standard monospace font.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterable, Sequence

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - developer convenience path
    raise SystemExit(
        "python-docx is required. Install it in your docs environment with "
        "`python3 -m pip install python-docx`."
    ) from exc


REPO_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = REPO_ROOT / "docs" / "google-docs"
SOURCES = (
    DOCS_DIR / "[INTERNAL] - Admin Toolkit Playbook.md",
    DOCS_DIR / "[INTERNAL] - Admin Toolkit Setup Guide.md",
)

BODY_FONT = "Times New Roman"
MONO_FONT = "Courier New"
INK = "172A30"
MUTED = "617177"
ACCENT = "168B8F"
ACCENT_DARK = "116B70"
LINE = "C7D1D4"
LIGHT = "EEF5F5"
CODE_BG = "F4F6F7"
TABLE_HEADER = "E7EFEF"


def _set_font(run_or_style, name: str, size: float | None = None) -> None:
    font = run_or_style.font
    font.name = name
    if size is not None:
        font.size = Pt(size)
    rpr = getattr(getattr(run_or_style, "_element", None), "rPr", None)
    if rpr is None:
        rpr = getattr(getattr(run_or_style, "_element", None), "get_or_add_rPr", lambda: None)()
    if rpr is not None:
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), name)


def _set_style_language(style, language: str = "en-US") -> None:
    rpr = style._element.get_or_add_rPr()
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), language)
    lang.set(qn("w:eastAsia"), language)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_font(normal, BODY_FONT, 12)
    _set_style_language(normal)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    _set_font(title, BODY_FONT, 24)
    _set_style_language(title)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("000000")
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True

    heading_specs = {
        "Heading 1": (23, "000000", 22, 8),
        "Heading 2": (17, ACCENT_DARK, 17, 5),
        "Heading 3": (13, "000000", 13, 3),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        _set_font(style, BODY_FONT, size)
        _set_style_language(style)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    # Google Docs imports these paragraph styles more predictably when their
    # typography is explicit instead of inherited from Word theme defaults.
    for name in ("List Bullet", "List Number"):
        if name in doc.styles:
            style = doc.styles[name]
            _set_font(style, BODY_FONT, 12)
            _set_style_language(style)
            style.font.color.rgb = RGBColor.from_string(INK)
            style.paragraph_format.space_after = Pt(3)

    if "ATK Code" not in doc.styles:
        code_style = doc.styles.add_style("ATK Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["ATK Code"]
    _set_font(code_style, MONO_FONT, 9)
    _set_style_language(code_style)
    code_style.font.color.rgb = RGBColor.from_string("26343A")
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0


def _set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str = LINE, size: str = "5") -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _prevent_row_split(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:cantSplit")) is None:
        trpr.append(OxmlElement("w:cantSplit"))


def _repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def _paragraph_border(paragraph, *, side: str, color: str, size: int, space: int = 4) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    pbdr.append(edge)


def _paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _add_hyperlink(paragraph, label: str, url: str, *, bold: bool = False, italic: bool = False) -> None:
    relation_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run_node = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    rfonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), BODY_FONT)
    rpr.append(rfonts)
    run_node.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = label
    run_node.append(text_node)
    hyperlink.append(run_node)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(
    r"(\[[^\]]+\]\(https?://[^)]+\)|\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*[^*]+\*(?!\*))"
)


def _add_inline(paragraph, text: str, *, force_bold: bool = False, force_italic: bool = False) -> None:
    """Add the deliberately small inline Markdown subset used by the sources."""
    pos = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.bold = force_bold
            run.italic = force_italic
            _set_font(run, BODY_FONT)
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            _add_hyperlink(paragraph, label, url, bold=force_bold, italic=force_italic)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.italic = force_italic
            _set_font(run, BODY_FONT)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.bold = force_bold
            run.italic = force_italic
            _set_font(run, MONO_FONT, 10.5)
            run.font.color.rgb = RGBColor.from_string("2F4A50")
        else:
            run = paragraph.add_run(token[1:-1])
            run.bold = force_bold
            run.italic = True
            _set_font(run, BODY_FONT)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = force_bold
        run.italic = force_italic
        _set_font(run, BODY_FONT)


def _set_paragraph_runs_color(paragraph, color: str) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    _paragraph_border(paragraph, side="bottom", color=LINE, size=6, space=1)


def _add_code_block(doc: Document, lines: Sequence[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    _set_table_borders(table, color="D9E0E2", size="4")
    cell = table.cell(0, 0)
    cell.width = Inches(6.45)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_shading(cell, CODE_BG)
    _set_cell_margins(cell, top=100, start=140, bottom=100, end=140)
    cell.text = ""
    first = True
    for line in lines or [""]:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        paragraph.style = doc.styles["ATK Code"]
        run = paragraph.add_run(line or " ")
        _set_font(run, MONO_FONT, 9)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_table(doc: Document, rows: Sequence[Sequence[str]], alignments: Sequence[str]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = doc.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)

    available = 6.5
    if column_count == 2:
        first_header = rows[0][0].strip().lower()
        first_width = 1.65 if first_header in {"attribute", "category", "decision", "mode", "resource"} else 2.15
        widths = (first_width, available - first_width)
    else:
        widths = tuple(available / column_count for _ in range(column_count))

    for idx, value in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(cell)
        _set_cell_shading(cell, TABLE_HEADER)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if alignments[idx] == "right" else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        _add_inline(paragraph, value, force_bold=True)
        _set_paragraph_runs_color(paragraph, INK)
    _repeat_table_header(table.rows[0])
    _prevent_row_split(table.rows[0])

    for row_values in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cell)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if alignments[idx] == "right" else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            _add_inline(paragraph, value)
            for run in paragraph.runs:
                if run.font.size is None:
                    run.font.size = Pt(10.5)
        _prevent_row_split(table.rows[-1])

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _parse_table(lines: Sequence[str], start: int) -> tuple[list[list[str]], list[str], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].lstrip().startswith("|"):
        values = [value.strip() for value in lines[idx].strip().strip("|").split("|")]
        rows.append(values)
        idx += 1
    separators = rows[1]
    alignments = ["right" if token.endswith(":") else "left" for token in separators]
    return [rows[0], *rows[2:]], alignments, idx


def _set_footer(doc: Document) -> None:
    for section in doc.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("INTERNAL ONLY  ·  Admin Toolkit  ·  v0.4.815  ·  ")
        _set_font(run, BODY_FONT, 8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        run = paragraph.add_run("Page ")
        _set_font(run, BODY_FONT, 8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        field_run = OxmlElement("w:r")
        field_rpr = OxmlElement("w:rPr")
        field_color = OxmlElement("w:color")
        field_color.set(qn("w:val"), MUTED)
        field_rpr.append(field_color)
        field_run.append(field_rpr)
        field_text = OxmlElement("w:t")
        field_text.text = "1"
        field_run.append(field_text)
        field.append(field_run)
        paragraph._p.append(field)


def _new_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    _configure_styles(doc)
    _set_footer(doc)
    return doc


def _add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    _add_inline(paragraph, text)
    if text.startswith("**Confidentiality Warning:**"):
        paragraph.paragraph_format.left_indent = Inches(0.16)
        paragraph.paragraph_format.right_indent = Inches(0.08)
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(9)
        _paragraph_border(paragraph, side="left", color=ACCENT, size=18, space=8)
        _paragraph_shading(paragraph, LIGHT)
    elif text == "**INTERNAL ONLY**":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_runs_color(paragraph, ACCENT_DARK)
        paragraph.paragraph_format.space_after = Pt(10)
    elif text.startswith("**Document baseline:**"):
        _set_paragraph_runs_color(paragraph, MUTED)
        for run in paragraph.runs:
            run.font.size = Pt(9.5)
        paragraph.paragraph_format.space_after = Pt(8)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.42)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(3)
    _add_inline(paragraph, text)


def _add_checklist(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(3)
    mark = paragraph.add_run("☐  ")
    _set_font(mark, "Arial Unicode MS", 11)
    mark.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    _add_inline(paragraph, text)


def _add_numbered(doc: Document, number: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.34)
    paragraph.paragraph_format.first_line_indent = Inches(-0.34)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{number}. ")
    run.bold = True
    _set_font(run, BODY_FONT)
    _add_inline(paragraph, text)


def _add_quote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.36)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(9)
    _paragraph_border(paragraph, side="left", color=ACCENT, size=16, space=7)
    _paragraph_shading(paragraph, "F5FAFA")
    _add_inline(paragraph, text, force_italic=True)


def render_markdown(source: Path, destination: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = _new_document()
    doc.core_properties.title = lines[0].lstrip("# ") if lines else source.stem
    doc.core_properties.subject = "Internal Admin Toolkit documentation"
    doc.core_properties.author = "Admin Toolkit maintainers"
    doc.core_properties.comments = f"Generated from {source.name}; optimized for Google Docs import."

    idx = 0
    in_code = False
    code_lines: list[str] = []
    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()

        if in_code:
            if stripped.startswith("```"):
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(raw)
            idx += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            idx += 1
            continue
        if not stripped:
            idx += 1
            continue
        if stripped == "---":
            _add_rule(doc)
            idx += 1
            continue
        if stripped.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].lstrip().startswith("|"):
            rows, alignments, idx = _parse_table(lines, idx)
            _add_table(doc, rows, alignments)
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text_value = heading.group(2)
            if level == 1:
                paragraph = doc.add_paragraph(style="Title")
                _add_inline(paragraph, text_value, force_bold=True)
            else:
                paragraph = doc.add_paragraph(style=f"Heading {level - 1}")
                _add_inline(paragraph, text_value, force_bold=True)
                if level == 2:
                    _paragraph_border(paragraph, side="bottom", color=LINE, size=5, space=4)
                # A long table is already a strong visual continuation. Do not
                # apply heading keep-with-next to the whole table: Pages treats
                # that as one oversized block and can leave a heading stranded
                # on an otherwise empty page during DOCX import.
                lookahead = idx + 1
                while lookahead < len(lines) and not lines[lookahead].strip():
                    lookahead += 1
                if lookahead < len(lines) and lines[lookahead].lstrip().startswith("|"):
                    table_end = lookahead
                    while table_end < len(lines) and lines[table_end].lstrip().startswith("|"):
                        table_end += 1
                    if table_end - lookahead >= 12:
                        paragraph.paragraph_format.keep_with_next = False
            idx += 1
            continue

        if stripped.startswith("> "):
            _add_quote(doc, stripped[2:].strip())
            idx += 1
            continue
        checklist = re.match(r"^-\s+\[\s\]\s+(.+)$", stripped)
        if checklist:
            _add_checklist(doc, checklist.group(1))
            idx += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            _add_bullet(doc, bullet.group(1))
            idx += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            _add_numbered(doc, numbered.group(1), numbered.group(2))
            idx += 1
            continue

        _add_body_paragraph(doc, stripped)
        idx += 1

    if in_code:
        _add_code_block(doc, code_lines)

    # Remove the empty starter paragraph Word creates only when it is actually
    # empty and precedes the title (normally add_paragraph reuses none, but this
    # keeps output deterministic across python-docx versions).
    body = doc._element.body
    first = body[0] if len(body) else None
    if first is not None and first.tag == qn("w:p") and not "".join(first.itertext()).strip():
        body.remove(first)

    doc.save(destination)


def main(argv: Iterable[str] | None = None) -> int:
    sources = list(SOURCES)
    missing = [str(path) for path in sources if not path.exists()]
    if missing:
        raise SystemExit("Missing documentation source(s):\n" + "\n".join(missing))
    for source in sources:
        destination = source.with_suffix(".docx")
        render_markdown(source, destination)
        print(f"generated {destination.relative_to(REPO_ROOT)}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
